"""
知识库管理模块
提供文档解析、文本分割、向量存储等功能
"""

import os
import uuid
import logging
from typing import Optional

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
)
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_text_splitters import Language
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# 代码文件扩展名到 LangChain Language 枚举的映射
CODE_LANGUAGE_MAP = {
    ".py": Language.PYTHON,
    ".js": Language.JS,
    ".ts": Language.TS,
    ".java": Language.JAVA,
    ".go": Language.GO,
    ".rs": Language.RUST,
    ".c": Language.C,
    ".cpp": Language.CPP,
    ".h": Language.CPP,
    ".hpp": Language.CPP,
    ".cs": Language.CSHARP,
    ".php": Language.PHP,
    ".rb": Language.RUBY,
    ".swift": Language.SWIFT,
    ".kt": Language.KOTLIN,
    ".scala": Language.SCALA,
    ".lua": Language.LUA,
    ".pl": Language.PERL,
    ".ps1": Language.POWERSHELL,
}


class KnowledgeBase:
    """知识库管理类，提供文档的增删查和向量检索功能"""

    def __init__(
        self,
        persist_directory: str = "./data/chroma_db",
        embedding_model_name: str = "qwen3-embedding:0.6b",
        base_url: str = "http://localhost:11434",
    ):
        """
        初始化知识库

        Args:
            persist_directory: Chroma 数据库持久化目录
            embedding_model_name: Ollama 嵌入模型名称（需使用专用的 embedding 模型，如 qwen3-embedding:0.6b）
            base_url: Ollama 服务地址
        """
        self.persist_directory = persist_directory
        self.embedding_model_name = embedding_model_name
        self.base_url = base_url

        # 确保存储目录存在
        os.makedirs(self.persist_directory, exist_ok=True)

        # 初始化嵌入模型
        self.embeddings = OllamaEmbeddings(
            model=self.embedding_model_name,
            base_url=self.base_url,
        )

        # 初始化向量数据库
        self.vectorstore = Chroma(
            persist_directory=self.persist_directory,
            embedding_function=self.embeddings,
        )

        # 初始化文本分割器-递归切分
        # 按段落→句子→字符的优先级递归切分，尽量在自然边界处切断，生产环境最常用的方案。
        # overlap的作用：相邻chunk之间重叠一部分文字，避免关键信息正好在切割点上被截断。
        # overlap 通常设 chunk\_size 的 10%-20%。
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
        )
        # 按token切分，留出余量，最大输入900token，低于模型1024上限
        # self.text_splitter = TokenTextSplitter(
        #     chunk_size=900,
        #     chunk_overlap=100,
        #     model_name="BAAI/bge-small-zh"
        # )

        logger.info("知识库初始化完成，存储目录: %s", self.persist_directory)

    def _parse_document(self, file_path: str) -> list[Document]:
        """
        根据文件类型解析文档

        Args:
            file_path: 文档文件路径

        Returns:
            解析后的文档列表

        Raises:
            ValueError: 不支持的文件类型
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == ".pdf":
            loader = PyPDFLoader(file_path)
        elif file_extension in (".txt", ".md"):
            loader = TextLoader(file_path, encoding="utf-8")
        elif file_extension == ".docx":
            # 使用 python-docx 直接解析
            docx_doc = DocxDocument(file_path)
            text = "\n".join(para.text for para in docx_doc.paragraphs)
            documents = [Document(page_content=text, metadata={"source": file_path})]
            logger.info("文档解析完成: %s, 共 1 段", file_path)
            return documents
        elif file_extension in CODE_LANGUAGE_MAP:
            # 代码文件按文本解析，切分时使用语言感知分割器
            loader = TextLoader(file_path, encoding="utf-8")
        else:
            raise ValueError(
                f"不支持的文件类型: {file_extension}。支持: .pdf, .txt, .md, .docx, "
                f".py, .js, .ts, .java, .go, .rs, .c, .cpp, .h, .cs, .php, .rb, "
                f".swift, .kt, .scala, .lua, .pl, .ps1"
            )

        documents = loader.load()
        logger.info("文档解析完成: %s, 共 %d 页/段", file_path, len(documents))
        return documents

    def _get_splitter(self, file_path: str):
        """
        根据文件类型获取合适的文本分割器

        代码文件使用语言感知分割器，按函数/类等语法结构切分，保留完整代码块；
        其他文件使用默认的递归字符分割器。

        Args:
            file_path: 文档文件路径

        Returns:
            文本分割器实例
        """
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension in CODE_LANGUAGE_MAP:
            language = CODE_LANGUAGE_MAP[file_extension]
            splitter = RecursiveCharacterTextSplitter.from_language(
                language=language,
                chunk_size=500,
                chunk_overlap=50,
            )
            logger.info("使用代码分割器: %s -> %s", file_extension, language)
            return splitter

        return self.text_splitter

    def add_document(self, file_path: str, file_name: Optional[str] = None) -> str:
        """
        添加文档到知识库

        Args:
            file_path: 文档文件路径
            file_name: 文档名称，默认为文件名

        Returns:
            文档唯一标识 ID
        """
        if file_name is None:
            file_name = os.path.basename(file_path)

        # 解析文档
        documents = self._parse_document(file_path)

        # 生成文档 ID
        doc_id = str(uuid.uuid4())

        # 为每个文档片段添加元数据
        for doc in documents:
            doc.metadata["doc_id"] = doc_id
            doc.metadata["file_name"] = file_name
            doc.metadata["file_path"] = file_path

        # 文本分割（代码文件使用语言感知分割器）
        splitter = self._get_splitter(file_path)
        split_docs = splitter.split_documents(documents)

        # 为分割后的片段添加元数据
        for i, doc in enumerate(split_docs):
            doc.metadata["doc_id"] = doc_id
            doc.metadata["file_name"] = file_name
            doc.metadata["file_path"] = file_path
            doc.metadata["chunk_index"] = i

        # 添加到向量数据库
        self.vectorstore.add_documents(split_docs)

        logger.info(
            "文档添加成功: %s (ID: %s), 分割为 %d 个文本块",
            file_name,
            doc_id,
            len(split_docs),
        )
        return doc_id

    def delete_document(self, doc_id: str) -> bool:
        """
        从知识库中删除文档

        Args:
            doc_id: 文档唯一标识 ID

        Returns:
            是否删除成功
        """
        try:
            # 获取向量数据库的底层 collection
            collection = self.vectorstore._collection

            # 根据 doc_id 元数据查找要删除的记录
            results = collection.get(where={"doc_id": doc_id})

            if not results["ids"]:
                logger.warning("未找到文档: %s", doc_id)
                return False

            # 删除匹配的记录
            collection.delete(ids=results["ids"])

            logger.info(
                "文档删除成功: %s, 共删除 %d 个文本块", doc_id, len(results["ids"])
            )
            return True
        except Exception as e:
            logger.error("删除文档失败: %s, 错误: %s", doc_id, str(e))
            raise

    def list_documents(self) -> list[dict]:
        """
        列出知识库中所有文档

        Returns:
            文档信息列表，每个元素包含 doc_id 和 file_name
        """
        try:
            # 获取所有记录
            collection = self.vectorstore._collection
            results = collection.get()

            # 根据 doc_id 去重，统计文档信息
            doc_map = {}
            for i, meta in enumerate(results.get("metadatas", [])):
                if meta and "doc_id" in meta:
                    doc_id = meta["doc_id"]
                    if doc_id not in doc_map:
                        doc_map[doc_id] = {
                            "doc_id": doc_id,
                            "file_name": meta.get("file_name", "未知"),
                            "file_path": meta.get("file_path", ""),
                            "chunk_count": 0,
                        }
                    doc_map[doc_id]["chunk_count"] += 1

            return list(doc_map.values())
        except Exception as e:
            logger.error("列出文档失败: %s", str(e))
            return []

    def get_retriever(self, search_kwargs: Optional[dict] = None):
        """
        获取检索器供 RAG 链使用

        Args:
            search_kwargs: 检索参数，如 {"k: 4}

        Returns:
            检索器对象
        """
        if search_kwargs is None:
            search_kwargs = {"k": 4}

        return self.vectorstore.as_retriever(search_kwargs=search_kwargs)
