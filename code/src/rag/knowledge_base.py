"""
知识库管理模块
提供文档解析、文本分割、向量存储等功能
"""

import os
import uuid
import logging
from typing import Optional

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
)
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_text_splitters import Language
from langchain_core.documents import Document

from llm_factory import LLMFactory
from config import (
    VECTOR_STORE_PROVIDER,
    MILVUS_URI,
    MILVUS_COLLECTION,
    CHROMA_PERSIST_DIRECTORY,
    CHROMA_COLLECTION,
)

logger = logging.getLogger(__name__)

# 代码文件扩展名到 LangChain Language 枚举的映射
CODE_LANGUAGE_MAP = {
    ".py": Language.PYTHON,
    ".js": Language.JS,
    ".ts": Language.TS,
    ".java": Language.JAVA,
    ".go": Language.GO,
    ".rs": Language.RUST,
    ".c": Language.C,
    ".cpp": Language.CPP,
    ".h": Language.CPP,
    ".hpp": Language.CPP,
    ".cs": Language.CSHARP,
    ".php": Language.PHP,
    ".rb": Language.RUBY,
    ".swift": Language.SWIFT,
    ".kt": Language.KOTLIN,
    ".scala": Language.SCALA,
    ".lua": Language.LUA,
    ".pl": Language.PERL,
    ".ps1": Language.POWERSHELL,
}


class KnowledgeBase:
    """知识库管理类，提供文档的增删查和向量检索功能"""

    def __init__(
        self,
        vector_store_provider: Optional[str] = None,
        milvus_uri: Optional[str] = None,
        collection_name: Optional[str] = None,
        embedding_model_name: Optional[str] = None,
        base_url: Optional[str] = None,
        provider: Optional[str] = None,
    ):
        """
        初始化知识库

        Args:
            vector_store_provider: 向量数据库提供商，为 None 时使用配置中的默认值。
                                   可选值: "milvus" | "chroma"
            milvus_uri: Milvus 服务地址，为 None 时使用配置中的默认值
            collection_name: 集合名称，为 None 时使用配置中的默认值
            embedding_model_name: 嵌入模型名称，为 None 时使用配置中的默认值
            base_url: 模型服务地址（保留用于向后兼容，实际使用 LLMFactory 中的配置）
            provider: 模型提供商，为 None 时使用配置中的默认值
        """
        self.vector_store_provider = vector_store_provider or VECTOR_STORE_PROVIDER
        self.milvus_uri = milvus_uri or MILVUS_URI
        self.collection_name = collection_name
        self.embedding_model_name = embedding_model_name
        self.base_url = base_url
        self.provider = provider

        # 初始化嵌入模型（使用 LLMFactory 支持多提供商）
        self.embeddings = LLMFactory.create_embedding(
            model_name=embedding_model_name,
            provider=provider,
        )

        # 初始化向量数据库
        self.vectorstore = self._create_vectorstore()

        # 初始化文本分割器
        self._init_text_splitter()

        logger.info("知识库初始化完成")

    def _create_vectorstore(self):
        """根据配置创建向量数据库实例"""
        if self.vector_store_provider == "milvus":
            from langchain_milvus import Milvus

            collection = self.collection_name or MILVUS_COLLECTION
            vectorstore = Milvus(
                connection_args={"uri": self.milvus_uri},
                collection_name=collection,
                embedding_function=self.embeddings,
            )
            logger.info(
                "Milvus 向量数据库初始化完成，地址: %s，集合: %s",
                self.milvus_uri,
                collection,
            )
            return vectorstore

        elif self.vector_store_provider == "chroma":
            from langchain_chroma import Chroma

            collection = self.collection_name or CHROMA_COLLECTION
            persist_dir = CHROMA_PERSIST_DIRECTORY or None
            vectorstore = Chroma(
                collection_name=collection,
                embedding_function=self.embeddings,
                persist_directory=persist_dir if persist_dir else None,
            )
            mode = "持久化模式" if persist_dir else "内存模式"
            logger.info(
                "Chroma 向量数据库初始化完成（%s），集合: %s",
                mode,
                collection,
            )
            return vectorstore

        else:
            raise ValueError(
                f"不支持的向量数据库提供商: {self.vector_store_provider}，可选值: milvus, chroma"
            )

    def _init_text_splitter(self):
        """初始化文本分割器"""
        logger.debug("初始化文本分割器: chunk_size=500, chunk_overlap=50")
        # 初始化文本分割器-递归切分
        # 按段落→句子→字符的优先级递归切分，尽量在自然边界处切断，生产环境最常用的方案。
        # overlap的作用：相邻chunk之间重叠一部分文字，避免关键信息正好在切割点上被截断。
        # overlap 通常设 chunk\_size 的 10%-20%。
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
        )

    def _parse_document(self, file_path: str) -> list[Document]:
        """
        根据文件类型解析文档

        Args:
            file_path: 文档文件路径

        Returns:
            解析后的文档列表

        Raises:
            ValueError: 不支持的文件类型
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == ".pdf":
            loader = PyPDFLoader(file_path)
        elif file_extension in (".txt", ".md"):
            loader = TextLoader(file_path, encoding="utf-8")
        elif file_extension == ".docx":
            # 使用 python-docx 直接解析
            docx_doc = DocxDocument(file_path)
            text = "\n".join(para.text for para in docx_doc.paragraphs)
            documents = [Document(page_content=text, metadata={"source": file_path})]
            logger.info("文档解析完成: %s, 共 1 段", file_path)
            return documents
        elif file_extension in CODE_LANGUAGE_MAP:
            # 代码文件按文本解析，切分时使用语言感知分割器
            loader = TextLoader(file_path, encoding="utf-8")
        else:
            raise ValueError(
                f"不支持的文件类型: {file_extension}。支持: .pdf, .txt, .md, .docx, "
                f".py, .js, .ts, .java, .go, .rs, .c, .cpp, .h, .cs, .php, .rb, "
                f".swift, .kt, .scala, .lua, .pl, .ps1"
            )

        documents = loader.load()
        logger.info("文档解析完成: %s, 共 %d 页/段", file_path, len(documents))
        return documents

    def _get_splitter(self, file_path: str):
        """
        根据文件类型获取合适的文本分割器

        代码文件使用语言感知分割器，按函数/类等语法结构切分，保留完整代码块；
        其他文件使用默认的递归字符分割器。

        Args:
            file_path: 文档文件路径

        Returns:
            文本分割器实例
        """
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension in CODE_LANGUAGE_MAP:
            language = CODE_LANGUAGE_MAP[file_extension]
            splitter = RecursiveCharacterTextSplitter.from_language(
                language=language,
                chunk_size=500,
                chunk_overlap=50,
            )
            logger.info("使用代码分割器: %s -> %s", file_extension, language)
            return splitter

        return self.text_splitter

    def add_document(self, file_path: str, file_name: Optional[str] = None) -> str:
        """
        添加文档到知识库

        Args:
            file_path: 文档文件路径
            file_name: 文档名称，默认为文件名

        Returns:
            文档唯一标识 ID
        """
        if file_name is None:
            file_name = os.path.basename(file_path)

        # 解析文档
        documents = self._parse_document(file_path)

        # 生成文档 ID
        doc_id = str(uuid.uuid4())

        # 为每个文档片段添加元数据
        for doc in documents:
            doc.metadata["doc_id"] = doc_id
            doc.metadata["file_name"] = file_name
            doc.metadata["file_path"] = file_path

        # 文本分割（代码文件使用语言感知分割器）
        splitter = self._get_splitter(file_path)
        split_docs = splitter.split_documents(documents)

        # 为分割后的片段添加元数据
        for i, doc in enumerate(split_docs):
            doc.metadata["doc_id"] = doc_id
            doc.metadata["file_name"] = file_name
            doc.metadata["file_path"] = file_path
            doc.metadata["chunk_index"] = i

        # 添加到向量数据库
        self.vectorstore.add_documents(split_docs)

        logger.info(
            "文档添加成功: %s (ID: %s), 分割为 %d 个文本块",
            file_name,
            doc_id,
            len(split_docs),
        )
        return doc_id

    def delete_document(self, doc_id: str) -> bool:
        """
        从知识库中删除文档

        Args:
            doc_id: 文档唯一标识 ID

        Returns:
            是否删除成功
        """
        try:
            if self.vector_store_provider == "milvus":
                # 使用底层 pymilvus client 删除匹配的记录
                client = self.vectorstore.client
                filter_expr = f'doc_id == "{doc_id}"'
                client.delete(
                    collection_name=self.collection_name,
                    filter=filter_expr,
                )
            else:
                # Chroma: 使用元数据过滤删除
                collection = self.vectorstore._collection
                results = collection.get(where={"doc_id": doc_id})
                if results["ids"]:
                    collection.delete(ids=results["ids"])

            logger.info("文档删除成功: %s", doc_id)
            return True
        except Exception as e:
            logger.error("删除文档失败: %s, 错误: %s", doc_id, str(e))
            raise

    def list_documents(self) -> list[dict]:
        """
        列出知识库中所有文档

        Returns:
            文档信息列表，每个元素包含 doc_id 和 file_name
        """
        logger.debug("开始列出知识库文档")
        try:
            if self.vector_store_provider == "milvus":
                # 使用底层 pymilvus client 直接查询元数据
                client = self.vectorstore.client
                collection = self.collection_name or MILVUS_COLLECTION
                results = client.query(
                    collection_name=collection,
                    filter='doc_id != ""',
                    output_fields=["doc_id", "file_name", "file_path"],
                    limit=10000,
                )
                doc_map = {}
                for row in results:
                    doc_id = row.get("doc_id", "")
                    if not doc_id:
                        continue
                    if doc_id not in doc_map:
                        doc_map[doc_id] = {
                            "doc_id": doc_id,
                            "file_name": row.get("file_name", "未知"),
                            "file_path": row.get("file_path", ""),
                            "chunk_count": 0,
                        }
                    doc_map[doc_id]["chunk_count"] += 1
            else:
                # Chroma: 使用底层 collection 查询
                collection = self.vectorstore._collection
                results = collection.get()
                doc_map = {}
                for i, meta in enumerate(results.get("metadatas", [])):
                    if meta and "doc_id" in meta:
                        doc_id = meta["doc_id"]
                        if doc_id not in doc_map:
                            doc_map[doc_id] = {
                                "doc_id": doc_id,
                                "file_name": meta.get("file_name", "未知"),
                                "file_path": meta.get("file_path", ""),
                                "chunk_count": 0,
                            }
                        doc_map[doc_id]["chunk_count"] += 1

            return list(doc_map.values())
        except Exception as e:
            logger.error("列出文档失败: %s", str(e))
            return []

    def get_retriever(self, search_kwargs: Optional[dict] = None):
        """
        获取检索器供 RAG 链使用

        Args:
            search_kwargs: 检索参数，如 {"k: 4}

        Returns:
            检索器对象
        """
        if search_kwargs is None:
            search_kwargs = {"k": 4}

        retriever = self.vectorstore.as_retriever(search_kwargs=search_kwargs)
        logger.debug("检索器已创建，参数: %s", search_kwargs)
        return retriever
